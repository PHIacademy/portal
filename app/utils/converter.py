from docx import Document
import os
import html

def convert_doc_to_html(file_obj):
    """Convert a DOC/DOCX file object to HTML format"""
    doc = Document(file_obj)
    html_content = []
    
    # Start with HTML structure
    html_content.append('<!DOCTYPE html>')
    html_content.append('<html><head>')
    html_content.append('<meta charset="UTF-8">')
    html_content.append('<style>')
    html_content.append('''
        body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }
        p { margin-bottom: 1em; }
        h1, h2, h3 { color: #333; margin-top: 1.5em; }
    ''')
    html_content.append('</style>')
    html_content.append('</head><body>')

    # Convert paragraphs
    for para in doc.paragraphs:
        if not para.text.strip():
            # Add empty paragraph for spacing
            html_content.append('<p>&nbsp;</p>')
            continue

        # Determine style
        if para.style.name.startswith('Heading'):
            level = para.style.name[-1]  # Get heading level
            html_content.append(f'<h{level}>{html.escape(para.text)}</h{level}>')
        else:
            # Process paragraph text with any inline formatting
            formatted_text = []
            for run in para.runs:
                text = html.escape(run.text)
                if run.bold:
                    text = f'<strong>{text}</strong>'
                if run.italic:
                    text = f'<em>{text}</em>'
                if run.underline:
                    text = f'<u>{text}</u>'
                formatted_text.append(text)
            
            html_content.append(f'<p>{"".join(formatted_text)}</p>')

    # Process tables if any
    for table in doc.tables:
        html_content.append('<table class="table table-bordered">')
        for row in table.rows:
            html_content.append('<tr>')
            for cell in row.cells:
                html_content.append(f'<td>{html.escape(cell.text)}</td>')
            html_content.append('</tr>')
        html_content.append('</table>')

    html_content.append('</body></html>')
    
    return '\n'.join(html_content)